# -*- coding: utf-8 -*-
"""Convertit PROCESS_LANCEMENT_STORES.md en .docx formaté (titres, tableaux, listes, gras, liens)."""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = sys.argv[1]
OUT = sys.argv[2]

doc = Document()

# Style de base
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

LINK_RE = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
BOLD_RE = re.compile(r'\*\*([^*]+)\*\*')
CODE_RE = re.compile(r'`([^`]+)`')


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1'); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_formatted(paragraph, text):
    """Ajoute du texte en gérant **gras**, `code`, [lien](url)."""
    # On découpe d'abord sur les liens
    pos = 0
    for m in LINK_RE.finditer(text):
        if m.start() > pos:
            _add_plain(paragraph, text[pos:m.start()])
        add_hyperlink(paragraph, m.group(2), m.group(1))
        pos = m.end()
    if pos < len(text):
        _add_plain(paragraph, text[pos:])


def _add_plain(paragraph, text):
    """Gère **gras** et `code` dans un segment sans lien."""
    # Tokenise gras et code
    tokens = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = paragraph.add_run(tok[2:-2]); run.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = 'Consolas'; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            paragraph.add_run(tok)


lines = open(SRC, encoding='utf-8').read().split('\n')
i = 0
while i < len(lines):
    line = lines[i]

    # Tableaux
    if line.strip().startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|[\s:|-]+\|\s*$', lines[i+1]):
        header = [c.strip() for c in line.strip().strip('|').split('|')]
        rows = []
        i += 2
        while i < len(lines) and lines[i].strip().startswith('|'):
            rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')])
            i += 1
        table = doc.add_table(rows=1, cols=len(header))
        table.style = 'Light Grid Accent 1'
        for j, h in enumerate(header):
            cell = table.rows[0].cells[j]
            cell.paragraphs[0].text = ''
            add_formatted(cell.paragraphs[0], h)
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for r in rows:
            cells = table.add_row().cells
            for j, c in enumerate(r):
                if j < len(cells):
                    cells[j].paragraphs[0].text = ''
                    add_formatted(cells[j].paragraphs[0], c)
        doc.add_paragraph()
        continue

    # Titres
    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=0)
    elif line.startswith('## '):
        doc.add_heading(re.sub(r'[#]', '', line[3:]).strip(), level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=2)
    elif line.startswith('#### '):
        doc.add_heading(line[5:].strip(), level=3)
    # Séparateur
    elif line.strip() == '---':
        pass
    # Citation
    elif line.startswith('> '):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
        add_formatted(p, line[2:].strip())
        for run in p.runs:
            run.italic = True
    # Bloc de code ```
    elif line.strip().startswith('```'):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i]); i += 1
        p = doc.add_paragraph()
        run = p.add_run('\n'.join(code_lines))
        run.font.name = 'Consolas'; run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.3)
    # Listes numérotées
    elif re.match(r'^\s*\d+\.\s', line):
        indent = len(line) - len(line.lstrip())
        p = doc.add_paragraph(style='List Number')
        if indent >= 2:
            p.paragraph_format.left_indent = Inches(0.5 + 0.25*(indent//2))
        add_formatted(p, re.sub(r'^\s*\d+\.\s', '', line))
    # Listes à puces
    elif re.match(r'^\s*[-*]\s', line):
        indent = len(line) - len(line.lstrip())
        p = doc.add_paragraph(style='List Bullet')
        if indent >= 2:
            p.paragraph_format.left_indent = Inches(0.5 + 0.25*(indent//2))
        add_formatted(p, re.sub(r'^\s*[-*]\s', '', line))
    # Paragraphe normal
    elif line.strip():
        p = doc.add_paragraph()
        add_formatted(p, line)

    i += 1

doc.save(OUT)
print('OK ->', OUT)
